import streamlit as st
import openai
from docx import Document
from docx.shared import Inches
from io import BytesIO
from PIL import Image
import tempfile
from fpdf import FPDF

# --- Streamlit page settings ---
st.set_page_config(page_title="Datasheet Generator", layout="centered")
st.title("📄 AI-Based Datasheet Generator")

# --- User input form ---
st.header("1. Product Input")

company_logo = st.file_uploader("Upload company logo (PNG or JPG)", type=["png", "jpg", "jpeg"])
product_name = st.text_input("Product Name")
description = st.text_area("Short Description")
sensor_type = st.selectbox("Sensor Type", ["Temperature", "Pressure", "Gas", "Accelerometer", "Other"])
dimensions = st.text_input("Dimensions (e.g., 50×30×10 mm)")
weight = st.text_input("Weight (g)")
power = st.text_input("Power Supply (e.g., 3.3 V / 150 mA)")
measurement_range = st.text_input("Measurement Range (e.g., -40°C to 125°C)")
sensitivity = st.text_input("Sensitivity (e.g., 10 mV/°C)")
accuracy = st.text_input("Accuracy (e.g., ±1°C)")
response_time = st.text_input("Response Time (e.g., <1s)")
signal_type = st.text_input("Output Signal (e.g., Analog / I2C / SPI)")
features = st.text_area("Key Features (comma separated)")
applications = st.text_area("Target Applications (comma separated)")

# --- Optional custom parameters ---
st.markdown("#### ➕ Optional Custom Parameters")
custom_fields = {}
custom_count = st.number_input("How many custom fields?", min_value=0, max_value=10, step=1)

for i in range(int(custom_count)):
    col1, col2 = st.columns(2)
    with col1:
        key_name = st.text_input(f"Custom Field Name #{i+1}", key=f"cf_k_{i}")
    with col2:
        key_value = st.text_input(f"Custom Field Value #{i+1}", key=f"cf_v_{i}")
    if key_name and key_value:
        custom_fields[key_name] = key_value

# --- Datasheet generation with GPT ---
if st.button("Generate Draft Datasheet"):
    with st.spinner("Generating content with GPT..."):
        # Compose prompt
        prompt = f"""
You are a professional technical writer specializing in sensor datasheets.
Generate a technical datasheet in clean markdown format, with the following structure and a section of "Recommended Additions" at the end:

1. Product Overview
2. Key Features
3. Applications
4. Absolute Maximum Ratings
5. Electrical Characteristics
6. Mechanical Dimensions
7. Output Interface
8. Additional Specifications
9. Recommended Additions

Respond in clean markdown. Do not guess missing values.

---
Input:
Name: {product_name}
Description: {description}
Sensor Type: {sensor_type}
Dimensions: {dimensions}
Weight: {weight}
Power: {power}
Measurement Range: {measurement_range}
Sensitivity: {sensitivity}
Accuracy: {accuracy}
Response Time: {response_time}
Output Signal: {signal_type}
Features: {features}
Applications: {applications}
"""
        for k, v in custom_fields.items():
            prompt += f"{k}: {v}\n"

        client = openai.OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a technical writer."},
                {"role": "user", "content": prompt}
            ]
        )
        datasheet_text = response.choices[0].message.content

        st.success("✅ Draft Datasheet Generated")
        st.markdown("### 📄 Preview")
        st.markdown(datasheet_text)

        # Load template and insert values
        doc = Document("Sensor_Datasheet_Template_Full.docx")
        for para in doc.paragraphs:
            if "<Product Name>" in para.text:
                para.text = para.text.replace("<Product Name>", product_name)
            if "<Short Description>" in para.text:
                para.text = para.text.replace("<Short Description>", description)

        for table in doc.tables:
            for row in table.rows[1:]:
                param = row.cells[0].text.strip().lower()
                if param == "dimensions" and dimensions:
                    row.cells[1].text = dimensions
                elif param == "weight" and weight:
                    row.cells[1].text = weight
                elif param == "power supply" and power:
                    row.cells[1].text = power
                elif param == "output signal" and signal_type:
                    row.cells[1].text = signal_type
                elif param == "measurement range" and measurement_range:
                    row.cells[1].text = measurement_range
                elif param == "sensitivity" and sensitivity:
                    row.cells[1].text = sensitivity
                elif param == "accuracy" and accuracy:
                    row.cells[1].text = accuracy
                elif param == "response time" and response_time:
                    row.cells[1].text = response_time
                else:
                    if not row.cells[1].text.strip():
                        row._element.getparent().remove(row._element)

        doc.add_heading("Recommended Additions", level=2)
        in_section = False
        for line in datasheet_text.splitlines():
            if line.strip().lower().startswith("### recommended additions"):
                in_section = True
                continue
            if in_section:
                if line.strip().startswith("-"):
                    doc.add_paragraph(line.strip("- ").strip(), style="List Bullet")
                elif line.strip() == "" or line.strip().startswith("###"):
                    break

        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        st.download_button("⬇️ Download Filled DOCX Draft", data=docx_buffer.getvalue(), file_name="sensor-datasheet-draft.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # --- PDF Export ---
        pdf = FPDF()
        pdf.add_page()
        if company_logo:
            logo_temp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            image = Image.open(company_logo).convert("RGB")
            image.save(logo_temp, format="PNG")
            logo_temp.flush()
            pdf.image(logo_temp.name, x=10, y=8, w=30)
            pdf.set_y(30)
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=11)
        for line in datasheet_text.splitlines():
            pdf.multi_cell(0, 10, line)
        pdf_buffer = BytesIO()
        pdf_output = pdf.output(dest='S').encode('latin-1')
        pdf_buffer.write(pdf_output)
        st.download_button("⬇️ Download as PDF", data=pdf_buffer.getvalue(), file_name="datasheet.pdf", mime="application/pdf")
